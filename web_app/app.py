import os
import io
import uuid
import zipfile
from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import hashlib
import json

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'

# Use environment variables for paths (Azure-compatible)
# Default to /tmp for Azure App Service which is writable
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', os.path.join(BASE_DIR, 'uploads'))
app.config['OUTPUT_FOLDER'] = os.getenv('OUTPUT_FOLDER', os.path.join(BASE_DIR, 'output'))
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024 * 1024  # 5GB max total upload size

# Ensure upload and output directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Log the paths for debugging
print(f"📁 UPLOAD_FOLDER: {app.config['UPLOAD_FOLDER']}")
print(f"📁 OUTPUT_FOLDER: {app.config['OUTPUT_FOLDER']}")
print(f"✅ Upload folder writable: {os.access(app.config['UPLOAD_FOLDER'], os.W_OK)}")
print(f"✅ Output folder writable: {os.access(app.config['OUTPUT_FOLDER'], os.W_OK)}")

ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_red_pdf_contents(pdf_path, original_filename=None):
    """Extract red content from PDF and return temp PDF path - V20 Logic"""
    print(f"\n{'='*60}")
    print(f"🔍 Starting red text extraction")
    print(f"📄 File: {original_filename or pdf_path}")
    print(f"📁 Source: {pdf_path}")
    print(f"📁 File exists: {os.path.exists(pdf_path)}")
    print(f"📊 File size: {os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 'N/A'} bytes")
    print(f"{'='*60}\n")
    
    doc = fitz.open(pdf_path)
    print(f"📖 PDF opened successfully. Pages: {doc.page_count}")
    
    # Configuration
    dpi_resol = 380
    font_size_flag = False
    
    # Create new PDF for red content
    new_pdf = fitz.open()
    red_content_found = False  # Track if any red content was found
    
    for page_num in range(doc.page_count):
        page_has_redtable = False
        tableposted_flag = False
        new_page = new_pdf.new_page(width=doc[page_num].rect.width, height=doc[page_num].rect.height)
        page_flag = False
        text_blocks = doc[page_num].get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
        flag_2 = False   # label level 2 outline is black or red
        flag_3 = False   # label level 3 outline is black or red
        page_figure_shown = False
        
        # Extract red content
        for i, block in enumerate(text_blocks):
            x0, y0, x1, y1, = block['bbox']
            rect = fitz.Rect(x0, y0, x1, y1)
            level_font_dict = set()
            block_flag = False
            flag_2_red_done = False 
            flag_3_red_done = False
            figure_flag = False
            figure_list = []
            block_printed = False
            
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"]
                    font_size = span.get("size", None)
                    font_name = span.get("font")
                    color = fitz.sRGB_to_rgb(span["color"])
                    is_bold = "Bold" in font_name or "Black" in font_name
                    level_font_dict.add((font_size, font_name))
                    
                    # Red color text detection
                    if color == (218, 31, 51):
                        if text == "•" or text == "●" or text == "∙" or text == "–":
                            continue
                        else:
                            if not page_flag:
                                print(f"🔴 Red text found on page {page_num + 1}: '{text[:50]}'")
                                page_flag = True
                            block_flag = True
                            red_content_found = True  # Mark that we found red content
                            
                            # Handle tables with red text
                            if ((font_size == 9 and font_name == "TimesNewRomanPSMT") or 
                                (font_size == 12 and font_name == "Arial-ItalicMT" and (text == "2024" or text == "2025"))) and (not tableposted_flag):
                                try:
                                    with pdfplumber.open(pdf_path) as pdf:
                                        for page_number, page in enumerate(pdf.pages, start=1):
                                            if page_number == page_num + 1:
                                                tables = page.find_tables()
                                                if not tables:
                                                    continue
                                                
                                                # Check each table for red text
                                                for table_index, table in enumerate(tables):
                                                    # Get the bounding box of the table
                                                    x0_t, top_t, x1_t, bottom_t = table.bbox
                                                    table_rect = fitz.Rect(x0_t, top_t, x1_t, bottom_t)
                                                    
                                                    # Check if this table contains red text
                                                    has_red_text = False
                                                    for blk in text_blocks:
                                                        blk_rect = fitz.Rect(blk['bbox'])
                                                        # Check if block overlaps with table
                                                        if table_rect.intersects(blk_rect):
                                                            if 'lines' in blk:
                                                                for ln in blk["lines"]:
                                                                    for sp in ln["spans"]:
                                                                        clr = fitz.sRGB_to_rgb(sp["color"])
                                                                        if clr == (218, 31, 51):
                                                                            has_red_text = True
                                                                            break
                                                                    if has_red_text:
                                                                        break
                                                            if has_red_text:
                                                                break
                                                    
                                                    # Only extract table if it contains red text
                                                    if has_red_text:
                                                        page_has_redtable = True
                                                        # Crop the table region from the page
                                                        pix = doc[page_num].get_pixmap(dpi=380, clip=(x0_t, top_t, x1_t, bottom_t))
                                                        
                                                        new_page.insert_image(fitz.Rect(x0_t, top_t, x1_t, bottom_t), pixmap=pix)
                                                
                                                tableposted_flag = True
                                except Exception as e:
                                    print(f"Error processing tables: {e}")
                    
                    elif color == (0, 0, 0) and not block_printed:
                        if is_bold == True and ((font_size == 14 and font_name == "Arial-Black") or 
                                               (font_size == 36 and font_name == "Arial-Black")):
                            image = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect))
                            new_page.insert_image(fitz.Rect(rect), pixmap=image)
                            block_printed = True
            
            # Detect level 2 and level 3 outline flags
            if block_flag:
                if (11.0, "Arial-BoldMT") in level_font_dict:
                    flag_2_red_done = True   # level 2 outline labeled by red or not
                elif (10.0, "Arial-BoldMT") in level_font_dict:
                    flag_3_red_done = True
            
            # Extract all font info from this block
            level_font_dict_all = set()
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"]
                    font_size = span.get("size", None)
                    font_name = span.get("font")
                    color = fitz.sRGB_to_rgb(span["color"])
                    is_bold = "Bold" in font_name or "Black" in font_name
                    level_font_dict_all.add((font_size, font_name))
                    
                    if font_size == 10:
                        font_size_flag = True
            
            # Update flags based on font detection
            if block_flag:
                if (11.0, "Arial-BoldMT") in level_font_dict_all:
                    flag_2, flag_3 = False, False
                    font_size_flag = False
                
                if (10.0, "Arial-BoldMT") in level_font_dict_all:
                    flag_3 = False
                
                if (12.0, "Arial-BoldMT") in level_font_dict_all:  # detect is figure related text or not
                    figure_flag = True
            
            image_2, image_3 = None, None
            rect_2, rect_3 = None, None
            
            # If goes to another chapter, then all level 2 and level 3 flag == 0
            if (12.0, "Arial-Black") in level_font_dict_all:
                flag_2, flag_3 = False, False
            else:  # Save level 2 and level 3 outline if not in red
                if (11.0, "Arial-BoldMT") in level_font_dict_all and not flag_2_red_done:
                    rect_2 = rect
                    image_2 = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect_2))
                    flag_2 = True
                
                if (10.0, "Arial-BoldMT") in level_font_dict_all and not flag_3_red_done:
                    rect_3 = rect
                    image_3 = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect_3))
                    flag_3 = True
            
            # Finally decide what exactly needs to be printed for this red content area
            if block_flag:
                # Print level 2 if not red color
                if not flag_2_red_done and flag_2 and rect_2 is not None:
                    new_page.insert_image(fitz.Rect(rect_2), pixmap=image_2)
                    rect_2 = None
                
                # Print level 3 if not red color
                if not flag_3_red_done and flag_3 and rect_3 is not None:
                    new_page.insert_image(fitz.Rect(rect_3), pixmap=image_3)
                    rect_3 = None
                
                # If figure title is red color, print out this figure
                if figure_flag:
                    if not page_figure_shown:
                        image = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect))
                        new_page.insert_image(fitz.Rect(rect), pixmap=image)
                        
                        print("has red figure here")
                        image_list = doc[page_num].get_images(full=True)
                        print(len(image_list))
                        
                        # Extract figures (excluding last one which is page outline)
                        for img_index, img in enumerate(image_list[:-1]):
                            xref = img[0]
                            rect_img = doc[page_num].get_image_bbox(img)
                            new_page.insert_image(fitz.Rect(rect_img), 
                                                pixmap=doc[page_num].get_pixmap(dpi=380, clip=rect_img))
                            page_figure_shown = True
                else:
                    # Don't print text inside tables, otherwise print all red text blocks
                    if page_has_redtable and (((9.0, "TimesNewRomanPS-BoldMT") in level_font_dict_all) or ((9.0, "TimesNewRomanPSMT") in level_font_dict_all)):
                        continue
                    else:
                        # Handle CambriaMath formulas with special vertical adjustment
                        if ((8.0, "CambriaMath") in level_font_dict_all) or ((10.5, "CambriaMath") in level_font_dict_all):
                            y_middle = (y0 + y1) / 2
                            y0_new = y_middle - 0.96 * (y1 - y_middle)
                            y1_new = y_middle + 0.96 * (y1 - y_middle)
                            rect = fitz.Rect(x0, y0_new, x1, y1_new)
                        
                        image = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect))
                        new_page.insert_image(fitz.Rect(rect), pixmap=image)
    
    # Save to temporary PDF file (required for image extraction)
    temp_pdf_path = os.path.join(app.config['OUTPUT_FOLDER'], f"temp_red_{uuid.uuid4().hex}.pdf")
    
    print(f"\n{'='*60}")
    if red_content_found:
        print(f"✅ Red content extraction SUCCESSFUL")
        print(f"📄 Created temp PDF with {len(new_pdf)} pages")
    else:
        print(f"⚠️  WARNING: No red text content found in PDF!")
        print(f"   This could mean:")
        print(f"   • The PDF doesn't contain red text")
        print(f"   • The red color doesn't match RGB(218, 31, 51)")
        print(f"   • The text is embedded as images")
    print(f"💾 Temp PDF: {temp_pdf_path}")
    print(f"{'='*60}\n")
    
    new_pdf.save(temp_pdf_path)
    new_pdf.close()
    doc.close()
    
    return temp_pdf_path

def extract_images_with_positions(pdf_path):
    """Extract images from PDF with their positions"""
    pdf_document = fitz.open(pdf_path)
    images = []
    
    print(f"Extracting images from PDF with {len(pdf_document)} pages...")
    
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        image_list = page.get_images(full=True)
        
        # Temporary list for images on this page
        page_images = []
        
        print(f"Page {page_num + 1}: Found {len(image_list)} images")
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            image_rects = page.get_image_rects(xref)
            for rect in image_rects:
                page_images.append({
                    "image_bytes": image_bytes,
                    "image_ext": image_ext,
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1,
                    "width": rect.width,
                    "height": rect.height,
                    "page_num": page_num
                })
                print(f"  Image {len(page_images)}: {rect.width}x{rect.height} at ({rect.x0}, {rect.y0})")
        
        # Sort images on this page by vertical position (y0)
        page_images.sort(key=lambda img: img["y0"])
        
        # Add sorted page images to the main list
        images.extend(page_images)
    
    print(f"Total images extracted: {len(images)}")
    pdf_document.close()
    
    # Delete temporary PDF after extraction
    try:
        os.remove(pdf_path)
        print(f"Deleted temporary PDF: {pdf_path}")
    except Exception as e:
        print(f"Error deleting temp PDF: {e}")
    
    return images

def create_word_document_with_positioned_images(images, output_docx_path):
    """Create Word document with positioned images"""
    doc = Document()
    
    print(f"Creating Word document with {len(images)} images...")
    
    if len(images) == 0:
        print("WARNING: No images found to add to Word document!")
        # Add a message to the document if no images found
        doc.add_paragraph("No red text content was found in the PDF file.")
        doc.add_paragraph("This could mean:")
        doc.add_paragraph("• The PDF doesn't contain red text")
        doc.add_paragraph("• The red color doesn't match the expected RGB values")
        doc.add_paragraph("• The text is embedded as images rather than text")
    else:
        for i, image_info in enumerate(images):
            try:
                image_bytes = image_info["image_bytes"]
                x0 = image_info["x0"]
                width = image_info["width"]
                height = image_info["height"]
                
                print(f"Adding image {i+1}: {width}x{height} at position {x0}")
                
                # Add paragraph for the image
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                
                # Add image to paragraph
                image_stream = BytesIO(image_bytes)
                # Limit image width to reasonable size (max 6 inches)
                max_width = min(width / 86, 6.0)
                run.add_picture(image_stream, width=Inches(max_width))
                
                # Set horizontal position
                paragraph.alignment = 0
                paragraph.paragraph_format.left_indent = Inches(x0 / 86)
                
            except Exception as e:
                print(f"Error adding image {i+1}: {e}")
    
    doc.save(output_docx_path)
    print(f"Word document saved: {output_docx_path}")
    return output_docx_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({'error': 'No files selected'}), 400
    
    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({'error': 'No files selected'}), 400
    
    job_id = str(uuid.uuid4())
    
    # Save uploaded files temporarily
    uploaded_files = []
    try:
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{job_id}_{filename}")
                file.save(temp_path)
                uploaded_files.append({
                    'filename': filename,
                    'path': temp_path
                })
        
        # Store job info in a simple dict (in production, use Redis or database)
        if not hasattr(app, 'jobs'):
            app.jobs = {}
        
        app.jobs[job_id] = {
            'files': uploaded_files,
            'total': len(uploaded_files),
            'processed': 0,
            'results': [],
            'status': 'pending'
        }
        
        return jsonify({
            'success': True,
            'job_id': job_id,
            'total_files': len(uploaded_files),
            'files': [f['filename'] for f in uploaded_files]
        })
    
    except Exception as e:
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/process/<job_id>', methods=['POST'])
def process_job(job_id):
    if not hasattr(app, 'jobs') or job_id not in app.jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = app.jobs[job_id]
    job['status'] = 'processing'
    
    try:
        for idx, file_info in enumerate(job['files']):
            filename = file_info['filename']
            temp_path = file_info['path']
            base_name = filename.rsplit('.', 1)[0]
            
            # Update current file being processed (1-based index for display)
            job['current_file'] = filename
            job['current_index'] = idx + 1  # 1-based index for user display
            
            # Extract red content from PDF file (returns temp PDF path)
            temp_red_pdf_path = extract_red_pdf_contents(temp_path, filename)
            
            # Extract images from processed PDF (deletes temp PDF after)
            images = extract_images_with_positions(temp_red_pdf_path)
            
            # Create Word document with simple filename
            # Use base_name directly, limit length for display
            if len(base_name) > 30:
                display_name = f"{base_name[:30]}.docx"
            else:
                display_name = f"{base_name}.docx"
            
            # Internal filename with job_id to avoid conflicts
            internal_filename = f"{job_id}_{display_name}"
            word_path = os.path.join(app.config['OUTPUT_FOLDER'], internal_filename)
            create_word_document_with_positioned_images(images, word_path)
            
            job['results'].append({
                'original_name': filename,
                'word_file': internal_filename,
                'display_name': display_name,  # Clean name for download
                'status': 'completed'
            })
            job['processed'] = idx + 1
            
            # Clean up temp file
            try:
                os.remove(temp_path)
            except:
                pass
        
        job['status'] = 'completed'
        
        return jsonify({
            'success': True,
            'job_id': job_id,
            'files': job['results'],
            'message': f'Successfully processed {len(job["results"])} files'
        })
    
    except Exception as e:
        job['status'] = 'failed'
        job['error'] = str(e)
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/progress/<job_id>')
def get_progress(job_id):
    if not hasattr(app, 'jobs') or job_id not in app.jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = app.jobs[job_id]
    return jsonify({
        'job_id': job_id,
        'status': job['status'],
        'total': job['total'],
        'processed': job['processed'],
        'current_index': job.get('current_index', 0),  # Which file is currently being processed
        'current_file': job.get('current_file', ''),
        'results': job['results']
    })

@app.route('/test_download')
def test_download():
    """Test route to check if downloads work"""
    try:
        output_files = os.listdir(app.config['OUTPUT_FOLDER'])
        if output_files:
            test_file = output_files[0]
            file_path = os.path.join(app.config['OUTPUT_FOLDER'], test_file)
            return send_file(file_path, as_attachment=True, download_name=test_file)
        else:
            return "No files in output folder"
    except Exception as e:
        return f"Error: {str(e)}"

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        # Ensure we have an absolute path
        file_path = os.path.abspath(file_path)
        
        print(f"Download request for: {filename}")
        print(f"Looking for file at: {file_path}")
        print(f"File exists: {os.path.exists(file_path)}")
        
        if os.path.exists(file_path):
            print(f"File size: {os.path.getsize(file_path)} bytes")
            
            # Extract clean display name (remove job_id prefix)
            display_name = filename
            if '_' in filename:
                parts = filename.split('_', 1)
                if len(parts) > 1:
                    display_name = parts[1]
            
            # Create response with proper headers
            response = send_file(
                file_path, 
                as_attachment=True, 
                download_name=display_name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            # Add additional security headers inline
            response.headers['X-Content-Type-Options'] = 'nosniff'
            response.headers['Content-Security-Policy'] = "default-src 'self'"
            
            return response
        else:
            print(f"File not found! Files in output folder:")
            if os.path.exists(app.config['OUTPUT_FOLDER']):
                for f in os.listdir(app.config['OUTPUT_FOLDER']):
                    print(f"  - {f}")
            else:
                print(f"Output folder doesn't exist: {app.config['OUTPUT_FOLDER']}")
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        print(f"Download error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_all/<job_id>')
def download_all(job_id):
    try:
        # Create zip file with all processed documents
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename in os.listdir(app.config['OUTPUT_FOLDER']):
                if filename.startswith(job_id):
                    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
                    # Remove job_id prefix from filename in zip
                    archive_name = filename.replace(f"{job_id}_", "")
                    zip_file.write(file_path, archive_name)
        
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'extracted_documents_{job_id}.zip'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/status/<job_id>')
def get_status(job_id):
    # Check status of processing job
    files_in_output = [f for f in os.listdir(app.config['OUTPUT_FOLDER']) if f.startswith(job_id)]
    return jsonify({
        'job_id': job_id,
        'status': 'completed' if files_in_output else 'processing',
        'files_count': len(files_in_output)
    })

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({
        'error': 'File(s) too large. Maximum total upload size is 5GB. Please upload fewer or smaller files.'
    }), 413

@app.after_request
def add_security_headers(response):
    """Add security headers to prevent download blocking"""
    # Only add headers for download routes
    if request.endpoint in ['download_file', 'download_all', 'test_download']:
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['Content-Disposition'] = response.headers.get('Content-Disposition', 'attachment')
        # Set Cache-Control to prevent caching issues
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

if __name__ == '__main__':
    print(f"\n{'='*60}")
    print(f"🚀 Application Starting")
    print(f"{'='*60}")
    print(f"📁 Upload Folder: {app.config['UPLOAD_FOLDER']}")
    print(f"📁 Output Folder: {app.config['OUTPUT_FOLDER']}")
    print(f"🌐 Server running at:")
    print(f"   - Local: http://127.0.0.1:5000")
    print(f"   - Network: http://10.240.34.164:5000")
    print(f"")
    print(f"💡 TIP: Use http://127.0.0.1:5000 for better browser compatibility")
    print(f"{'='*60}\n")
    app.run(debug=True, host='0.0.0.0', port=5000)