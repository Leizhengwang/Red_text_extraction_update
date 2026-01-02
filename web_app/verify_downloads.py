#!/usr/bin/env python3
"""
Verify that downloaded Word files are valid and readable
"""

import os
from docx import Document

def verify_word_file(file_path):
    """Verify that a Word file is valid and can be opened"""
    try:
        doc = Document(file_path)
        paragraphs = len(doc.paragraphs)
        images = len(doc.inline_shapes)
        
        print(f"✅ {os.path.basename(file_path)}:")
        print(f"   - File size: {os.path.getsize(file_path):,} bytes")
        print(f"   - Paragraphs: {paragraphs}")
        print(f"   - Inline shapes/images: {images}")
        
        # Check if there's any text content
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
        
        if text_content:
            print(f"   - Text content preview: {text_content[0][:100]}...")
        else:
            print(f"   - No text content found")
        
        return True
        
    except Exception as e:
        print(f"❌ {os.path.basename(file_path)}: Error - {e}")
        return False

def main():
    output_dir = "output"
    print("Verifying Word document downloads...\n")
    
    if not os.path.exists(output_dir):
        print(f"❌ Output directory {output_dir} not found")
        return
    
    word_files = [f for f in os.listdir(output_dir) if f.endswith('.docx')]
    
    if not word_files:
        print("❌ No Word files found in output directory")
        return
    
    print(f"Found {len(word_files)} Word files to verify:\n")
    
    valid_files = 0
    for filename in word_files:
        file_path = os.path.join(output_dir, filename)
        if verify_word_file(file_path):
            valid_files += 1
        print()
    
    print(f"Summary: {valid_files}/{len(word_files)} files are valid Word documents")

if __name__ == "__main__":
    main()